import json
import os
import shutil
from datetime import datetime, timezone, timedelta
import pandas as pd
import streamlit as st
import urllib.parse
from docx import Document
import re
import streamlit.components.v1 as components

try:
    import cv2
    import numpy as np
    OPENCV_DISPONIVEL = True
except ImportError:
    OPENCV_DISPONIVEL = False

st.set_page_config(
    page_title="Premium Wines - Galpão",
    page_icon="imagem premium.jpeg",
    layout="wide",
    initial_sidebar_state="collapsed",
)

st.markdown(
    """
    <style>
    .stApp { background: linear-gradient(135deg, #F8F9FA 0%, #E9ECEF 100%); color: #1A1A1A; font-family: 'Poppins', sans-serif; overscroll-behavior-y: none; }
    [data-testid="stSidebar"] { display: none; }
    
    footer {visibility: hidden;}
    #MainMenu {visibility: hidden;}
    [data-testid="stStatusWidget"] {display: none;}
    
    label { color: #7A1C2E !important; font-weight: 700 !important; font-size: 0.95rem !important; }
    .wine-card { background-color: #FFFFFF; color: #1A1A1A; border-radius: 14px; padding: 16px; margin-bottom: 12px; border: 1px solid #E9ECEF; box-shadow: 0px 4px 12px rgba(0, 0, 0, 0.03); }
    .wine-title { color: #7A1C2E; font-size: 1.1rem; font-weight: 700; margin-bottom: 4px; }
    .badge-pallet-grande { background-color: #7A1C2E; color: #FFFFFF; padding: 6px 14px; border-radius: 8px; font-weight: 700; font-size: 1rem; display: inline-block; }
    .stButton button { background-color: #7A1C2E !important; color: #FFFFFF !important; border-radius: 12px !important; font-weight: 600 !important; border: none !important; padding: 10px 16px !important; width: 100%; white-space: pre-wrap; }
    </style>
""", unsafe_allow_html=True,
)

NOME_ARQUIVO = "estoque_galpao_pro.json"
ARQUIVO_USUARIOS = "usuarios_galpao.json"
ARQUIVO_LOGS = "logs_auditoria.json"
ARQUIVO_PEDIDOS = "pedidos_matriz.json"
PASTA_BACKUP = "backups_estoque"
PASTA_FOTOS = "fotos_vinhos"
SENHA_DEV = "1980"

if not os.path.exists(PASTA_BACKUP):
    os.makedirs(PASTA_BACKUP)
if not os.path.exists(PASTA_FOTOS):
    os.makedirs(PASTA_FOTOS)

LISTA_CORREDORES = [f"Corredor {i:02d}" for i in range(1, 26)]
LISTA_LOCAIS_TIPO = ["Pallet", "Prateleira"]
LISTA_NUMEROS_LOCAL = [f"Item {i:02d}" for i in range(1, 26)]
LISTA_LADOS = ["Direito", "Esquerdo", "Centro / Único"]
OPCOES_CAIXA = ["Caixa com 12 garrafas", "Caixa com 6 garrafas", "Caixa com 3 garrafas", "Caixa com 2 garrafas", "Garrafa Avulsa (1 un)", "Outra quantidade"]

def obter_saudacao():
    fuso = timezone(timedelta(hours=-3))
    hora = datetime.now(fuso).hour
    if 0 <= hora < 12: return "Bom dia"
    elif 12 <= hora < 18: return "Boa tarde"
    else: return "Boa noite"

def realizar_backup(nome):
    if os.path.exists(nome):
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        shutil.copy(nome, os.path.join(PASTA_BACKUP, f"backup_{ts}_{nome}"))

def carregar_dados():
    estoque = []
    if os.path.exists(NOME_ARQUIVO):
        try:
            with open(NOME_ARQUIVO, "r", encoding="utf-8") as f: estoque = json.load(f)
        except: pass
    if not estoque:
        estoque = [{"nome": "Campana Merlot", "tipo": "Tinto", "safra": "2024", "localizacao": "Corredor 01 - Pallet Item 01", "lado": "Direito", "caixa": "Caixa com 12 garrafas", "codigo_barras": "7891008116632", "foto": ""}]
    return sorted(estoque, key=lambda x: x.get("nome", "").lower())

def salvar_dados(estoque):
    estoque_ordenado = sorted(estoque, key=lambda x: x.get("nome", "").lower())
    with open(NOME_ARQUIVO, "w", encoding="utf-8") as f: json.dump(estoque_ordenado, f, ensure_ascii=False, indent=4)
    realizar_backup(NOME_ARQUIVO)

def carregar_usuarios():
    if os.path.exists(ARQUIVO_USUARIOS):
        try:
            with open(ARQUIVO_USUARIOS, "r", encoding="utf-8") as f: return json.load(f)
        except: pass
    return [{"nome": "Vagner Souza", "cargo": "Administrador", "senha": "1980"}]

def salvar_usuarios(usuarios):
    with open(ARQUIVO_USUARIOS, "w", encoding="utf-8") as f: json.dump(usuarios, f, ensure_ascii=False, indent=4)

def carregar_logs():
    if os.path.exists(ARQUIVO_LOGS):
        try:
            with open(ARQUIVO_LOGS, "r", encoding="utf-8") as f: return json.load(f)
        except: pass
    return []

def registrar_log(usuario, acao, detalhes):
    logs = carregar_logs()
    logs.insert(0, {"data_hora": datetime.now().strftime("%d/%m/%Y %H:%M:%S"), "usuario": usuario, "acao": acao, "detalhes": detalhes})
    with open(ARQUIVO_LOGS, "w", encoding="utf-8") as f: json.dump(logs, f, ensure_ascii=False, indent=4)

def carregar_pedidos():
    if os.path.exists(ARQUIVO_PEDIDOS):
        try:
            with open(ARQUIVO_PEDIDOS, "r", encoding="utf-8") as f: return json.load(f)
        except: pass
    return []

def salvar_pedidos(pedidos):
    with open(ARQUIVO_PEDIDOS, "w", encoding="utf-8") as f: json.dump(pedidos, f, ensure_ascii=False, indent=4)

def gerar_qr_code_api(texto):
    return f"https://api.qrserver.com/v1/create-qr-code/?size=250x250&data={urllib.parse.quote(texto)}"

def interpretar_linha_pedido(texto_linha):
    texto = texto_linha.strip()
    safra = ""
    quantidade = 1
    
    anos = re.findall(r'\b(20\d{2})\b', texto)
    if anos:
        safra = anos[0]
        texto_limpo = texto.replace(safra, "")
    else:
        texto_limpo = texto

    match_qtd = re.search(r'(?:/|\bcaixas?|\bqt[d]?\.?)\s*(\d+)', texto_limpo, re.IGNORECASE)
    if match_qtd:
        quantidade = int(match_qtd.group(1))
        texto_limpo = texto_limpo.replace(match_qtd.group(0), "")
    else:
        numeros_soltos = re.findall(r'\b(\d+)\b', texto_limpo)
        if numeros_soltos and numeros_soltos[-1] != safra:
            quantidade = int(numeros_soltos[-1])
            texto_limpo = texto_limpo.replace(numeros_soltos[-1], "")

    texto_limpo = re.sub(r'\bcaixas?\b', '', texto_limpo, flags=re.IGNORECASE)
    nome = re.sub(r'[/\|\-\–]+', '', texto_limpo).strip().title()
    return {"nome": nome, "safra": safra, "quantidade": quantidade, "separado": False, "qtd_separada": 0}

def extrair_pedidos_de_arquivo(arq):
    itens = []
    ext = arq.name.split('.')[-1].lower()
    try:
        if ext in ['xlsx', 'xls']:
            df = pd.read_excel(arq)
            for _, row in df.iterrows():
                nome_bruto = str(row.get('Nome', row.iloc[0] if len(row) > 0 else '')).strip()
                if nome_bruto and nome_bruto != 'Nan':
                    safra_col = str(row.get('Safra', row.iloc[1] if len(row) > 1 else '')).strip()
                    qtd_col = row.get('Quantidade', row.iloc[2] if len(row) > 2 else 1)
                    try: qtd = int(qtd_col) if pd.notnull(qtd_col) else 1
                    except: qtd = 1
                    
                    itens.append({
                        "nome": nome_bruto.title(), 
                        "safra": safra_col if safra_col != 'Nan' else '', 
                        "quantidade": qtd, 
                        "separado": False, 
                        "qtd_separada": 0
                    })
        elif ext == 'txt':
            linhas = [l.strip() for l in arq.getvalue().decode("utf-8").split("\n") if l.strip()]
            for l in linhas:
                itens.append(interpretar_linha_pedido(l))
    except: pass
    return itens

def extrair_linhas_de_arquivo(arq):
    linhas = []
    ext = arq.name.split('.')[-1].lower()
    try:
        if ext in ['xlsx', 'xls']:
            df = pd.read_excel(arq)
            for c in df.columns:
                for v in df[c].dropna():
                    if str(v).strip(): linhas.append(str(v).strip().title())
        elif ext == 'docx':
            doc = Document(arq)
            for p in doc.paragraphs:
                if p.text.strip(): linhas.append(p.text.strip().title())
            for t in doc.tables:
                for row in t.rows:
                    for cell in row.cells:
                        if cell.text.strip(): linhas.append(cell.text.strip().title())
        elif ext == 'txt':
            linhas = [l.strip().title() for l in arq.getvalue().decode("utf-8").split("\n") if l.strip()]
    except: pass
    return linhas

def extrair_numero_corredor(localizacao):
    nums = re.findall(r'\d+', localizacao)
    if nums:
        return int(nums[0])
    return 999

def componente_leitor_barcode(chave_sessao):
    codigo_atual = st.session_state.get(chave_sessao, "")
    
    html_code = f"""
    <div style="text-align: center;">
        <div id="reader_{chave_sessao}" style="width: 100%; max-width: 400px; margin: auto; border-radius: 12px; overflow: hidden;"></div>
        <p id="resultado_{chave_sessao}" style="font-weight: bold; color: #7A1C2E; margin-top: 8px;"></p>
    </div>
    
    <script src="https://unpkg.com/html5-qrcode"></script>
    <script>
      function onScanSuccess(decodedText, decodedResult) {{
        document.getElementById("resultado_{chave_sessao}").innerText = "Lido com sucesso: " + decodedText;
        
        const url = new URL(window.parent.location.href);
        url.searchParams.set('scanned_{chave_sessao}', decodedText);
        window.parent.history.replaceState({{}}, '', url);
        
        const triggerRerun = new Event('input', {{ bubbles: true }});
        
        if (window.html5QrCode_{chave_sessao}) {{
            window.html5QrCode_{chave_sessao}.stop().catch(err => {{}});
        }}
      }}

      try {{
          const html5QrCode = new Html5Qrcode("reader_{chave_sessao}");
          window.html5QrCode_{chave_sessao} = html5QrCode;
          
          const config = {{ fps: 10, qrbox: {{ width: 250, height: 150 }} }};
          html5QrCode.start({{ facingMode: "environment" }}, config, onScanSuccess).catch(err => {{
              document.getElementById("resultado_{chave_sessao}").innerText = "Erro ao acessar câmera: verifique permissões.";
          }});
      }} catch (e) {{}}
    </script>
    """
    components.html(html_code, height=320)

if "usuarios" not in st.session_state:
    st.session_state.usuarios = carregar_usuarios()

st.session_state.estoque = carregar_dados()
st.session_state.pedidos = carregar_pedidos()

if "menu_atual" not in st.session_state: st.session_state.menu_atual = "🏠 Home"
if "termo_busca" not in st.session_state: st.session_state.termo_busca = ""
if "codigo_capturado_cadastro" not in st.session_state: st.session_state.codigo_capturado_cadastro = ""
if "codigos_bipados_conferencia" not in st.session_state: st.session_state.codigos_bipados_conferencia = {}

qp = st.query_params

for key, val in qp.items():
    if key.startswith("scanned_"):
        sess_key = key.replace("scanned_", "")
        if sess_key == "cad_barcode":
            st.session_state.codigo_capturado_cadastro = val
        else:
            st.session_state.codigos_bipados_conferencia[sess_key] = val
        del st.query_params[key]
        st.rerun()

user_url = qp.get("user", None)
cargo_url = qp.get("cargo", "Operador")

if "usuario_logado" not in st.session_state or st.session_state.usuario_logado is None:
    if user_url:
        st.session_state.usuario_logado = {"nome": user_url, "cargo": cargo_url}
    else:
        st.session_state.usuario_logado = None

if st.session_state.usuario_logado is None:
    st.write("")
    _, cc, _ = st.columns([1, 1.3, 1])
    with cc:
        if os.path.exists("imagem premium.jpeg"):
            _, ci, _ = st.columns([1, 1.8, 1])
            with ci: st.image("imagem premium.jpeg", width=190)
        st.markdown("<h1 style='text-align: center; color: #7A1C2E; font-size: 1.6rem;'>PREMIUM WINES GALPÃO</h1>", unsafe_allow_html=True)
        
        tab1, tab2, tab3 = st.tabs(["🔑 Entrar", "👤 Criar Conta", "⚙️ Dev"])
        with tab1:
            with st.form("l_form"):
                u = st.text_input("Usuário").strip().title()
                p = st.text_input("Senha", type="password").strip()
                if st.form_submit_button("ENTRAR", use_container_width=True):
                    user = next((x for x in st.session_state.usuarios if x['nome'].lower() == u.lower() and x['senha'] == p), None)
                    if user:
                        st.session_state.usuario_logado = user
                        st.query_params["user"] = user['nome']
                        st.query_params["cargo"] = user.get('cargo', 'Operador')
                        st.rerun()
                    else: st.error("Dados incorretos.")
        with tab2:
            with st.form("c_form"):
                n = st.text_input("Nome").strip().title()
                s = st.text_input("Senha", type="password").strip()
                if st.form_submit_button("CADASTRAR", use_container_width=True):
                    if n and s:
                        novo = {"nome": n, "cargo": "Operador", "senha": s}
                        st.session_state.usuarios.append(novo)
                        salvar_usuarios(st.session_state.usuarios)
                        st.session_state.usuario_logado = novo
                        st.query_params["user"] = novo['nome']
                        st.query_params["cargo"] = novo['cargo']
                        st.rerun()
                    else: st.error("Preencha tudo.")
        with tab3:
            with st.form("d_form"):
                sp = st.text_input("Senha Mestra", type="password")
                if st.form_submit_button("DEV", use_container_width=True):
                    if sp == SENHA_DEV:
                        dev_user = {"nome": "Dev", "cargo": "Desenvolvedor"}
                        st.session_state.usuario_logado = dev_user
                        st.query_params["user"] = "Dev"
                        st.query_params["cargo"] = "Desenvolvedor"
                        st.rerun()
                    else: st.error("Senha incorreta.")
    st.stop()

ct1, ct2, ct3 = st.columns([3, 2, 1])
with ct1: st.markdown(f"🍷 <b>PREMIUM WINES</b> | Usuário: {st.session_state.usuario_logado['nome']} ({st.session_state.usuario_logado.get('cargo', 'Operador')})", unsafe_allow_html=True)
with ct2:
    if st.session_state.menu_atual != "🏠 Home":
        if st.button("⬅️ Voltar ao Menu", use_container_width=True): st.session_state.menu_atual = "🏠 Home"; st.rerun()
with ct3:
    if st.button("🚪 Sair", use_container_width=True):
        st.session_state.usuario_logado = None
        st.query_params.clear()
        st.session_state.menu_atual = "🏠 Home"
        st.rerun()

st.markdown("---")

if st.session_state.menu_atual == "🏠 Home":
    st.markdown(f"<p style='text-align: center; color: #666; margin-bottom: 0px;'>{obter_saudacao()},</p>", unsafe_allow_html=True)
    st.markdown(f"<h1 style='text-align: center; color: #7A1C2E; margin-top: 0px;'>{st.session_state.usuario_logado['nome']}! 👋</h1>", unsafe_allow_html=True)
    st.markdown("<p style='text-align: center; color: #444; font-size: 0.95rem; margin-bottom: 25px;'>Escolha abaixo a opção desejada para gerenciar o galpão:</p>", unsafe_allow_html=True)
    
    c1, c2, c3 = st.columns(3)
    with c1:
        if st.button("📦 Lista de Pedidos (Matriz)", use_container_width=True): st.session_state.menu_atual = "PedidosMatriz"; st.rerun()
    with c2:
        if st.button("🔍 Buscar / Filtros", use_container_width=True): st.session_state.menu_atual = "Filtros"; st.rerun()
    with c3:
        if st.button("🗺️ Mapa de Separação", use_container_width=True): st.session_state.menu_atual = "MapaSeparacao"; st.rerun()
    
    st.write("")
    c4, c5, c6 = st.columns(3)
    with c4:
        if st.button("🍷 Estoque Completo", use_container_width=True): st.session_state.menu_atual = "Estoque"; st.rerun()
    with c5:
        if st.button("➕ Cadastrar Vinho", use_container_width=True): st.session_state.menu_atual = "Cadastrar"; st.rerun()
    with c6:
        if st.button("📱 Gerar QR Code", use_container_width=True): st.session_state.menu_atual = "GerarQR"; st.rerun()
        
    st.write("")
    c7, c8, c9 = st.columns(3)
    with c7:
        if st.button("📷 Escanear Local", use_container_width=True): st.session_state.menu_atual = "Scanner"; st.rerun()
    with c8:
        if st.button("✏️ Editar Vinho", use_container_width=True): st.session_state.menu_atual = "Editar"; st.rerun()
    with c9:
        if st.button("📋 Histórico", use_container_width=True): st.session_state.menu_atual = "Historico"; st.rerun()
        
    if st.session_state.usuario_logado.get('cargo') == "Desenvolvedor":
        st.write("")
        if st.button("⚙️ Gerenciar Contas", use_container_width=True):
            st.session_state.menu_atual = "GerenciarUsuarios"
            st.rerun()

elif st.session_state.menu_atual == "PedidosMatriz":
    st.subheader("📦 Gerenciamento de Pedidos e Conferência Antierros")
    
    aba_ped1, aba_ped2 = st.tabs(["📋 Enviar/Novo Pedido", "🔍 Roteiro e Leitor de Caixa (Antierros)"])
    
    with aba_ped1:
        st.markdown("Envie a lista enviada pela matriz (Excel ou TXT) ou digite livremente (ex: *Campana Merlot 2024 /05 Caixas*).")
        
        with st.form("form_novo_pedido"):
            id_pedido = st.text_input("Identificação do Pedido / Loja", value=f"Pedido #{datetime.now().strftime('%d/%m %H:%M')}")
            arq_pedido = st.file_uploader("Arquivo de Pedido (Excel ou TXT)", type=["xlsx", "xls", "txt"])
            
            st.markdown("Ou digite os itens linha a linha (o sistema lê automaticamente a safra e a quantidade):")
            texto_manual_pedido = st.text_area("Ex: Campana Merlot 2024 /05 Caixas", placeholder="Campana Merlot 2024 / 05 Caixas\nLa Consulta Malbec 2023 / 2")
            
            if st.form_submit_button("Cadastrar Pedido para Separação"):
                itens_novos = []
                if arq_pedido is not None:
                    itens_novos = extrair_pedidos_de_arquivo(arq_pedido)
                if texto_manual_pedido.strip():
                    for linha in texto_manual_pedido.split("\n"):
                        if linha.strip():
                            itens_novos.append(interpretar_linha_pedido(linha))
                
                if itens_novos:
                    novo_registro_pedido = {
                        "id": id_pedido,
                        "data": datetime.now().strftime("%d/%m/%Y %H:%M"),
                        "itens": itens_novos,
                        "status": "Pendente"
                    }
                    st.session_state.pedidos.append(novo_registro_pedido)
                    salvar_pedidos(st.session_state.pedidos)
                    registrar_log(st.session_state.usuario_logado['nome'], "Novo Pedido Matriz", id_pedido)
                    st.success("Pedido cadastrado com sucesso!")
                    st.rerun()
                else:
                    st.error("Adicione itens por arquivo ou texto.")
                    
    with aba_ped2:
        if not st.session_state.pedidos:
            st.info("Nenhum pedido cadastrado.")
        else:
            pedidos_nomes = [f"{p['id']} ({p['data']}) - Status: {p.get('status', 'Pendente')}" for p in st.session_state.pedidos]
            escolha_ped_str = st.selectbox("Selecione o Pedido:", pedidos_nomes)
            idx_ped = pedidos_nomes.index(escolha_ped_str)
            pedido_atual = st.session_state.pedidos[idx_ped]
            
            col_info_ped, col_del_ped = st.columns([3, 1])
            with col_info_ped:
                st.markdown(f"### Pedido: {pedido_atual['id']}")
            with col_del_ped:
                if st.button("🗑️ Excluir Pedido", use_container_width=True):
                    id_removido = pedido_atual['id']
                    st.session_state.pedidos.pop(idx_ped)
                    salvar_pedidos(st.session_state.pedidos)
                    registrar_log(st.session_state.usuario_logado['nome'], "Excluir Pedido", id_removido)
                    st.success("Pedido excluído com sucesso!")
                    st.rerun()
            
            ordem_separacao = st.radio("Direção da Rota pelos Corredores:", ["Crescente (Corredor 01 ao 25)", "Decrescente (Corredor 25 ao 01)"], horizontal=True)
            reversa = True if "Decrescente" in ordem_separacao else False
            
            itens_com_local = []
            for item in pedido_atual['itens']:
                nome_pedido_limpo = item['nome'].lower()
                safra_pedido = str(item.get('safra', '')).strip()
                
                vinho_encontrado = None
                for v in st.session_state.estoque:
                    v_nome = v.get('nome', '').lower()
                    v_safra = str(v.get('safra', '')).strip()
                    
                    palavras_pedido = [p for p in nome_pedido_limpo.split() if len(p) > 2]
                    match_parcial = any(p in v_nome for p in palavras_pedido) if palavras_pedido else False
                    
                    if v_nome in nome_pedido_limpo or nome_pedido_limpo in v_nome or match_parcial:
                        if safra_pedido and v_safra:
                            if safra_pedido == v_safra:
                                vinho_encontrado = v
                                break
                        else:
                            vinho_encontrado = v
                            break
                            
                if not vinho_encontrado:
                    vinho_encontrado = next((v for v in st.session_state.estoque if v.get('nome', '').lower() in nome_pedido_limpo or nome_pedido_limpo in v.get('nome', '').lower()), None)
                
                loc = vinho_encontrado.get('localizacao', 'Corredor 01 - Pallet Item 01') if vinho_encontrado else 'Não cadastrado'
                corr_num = extrair_numero_corredor(loc)
                
                itens_com_local.append({
                    "item_original": item,
                    "vinho_estoque": vinho_encontrado,
                    "localizacao": loc,
                    "corredor_num": corr_num
                })
                
            itens_ordenados = sorted(itens_com_local, key=lambda x: x['corredor_num'], reverse=reversa)
            
            st.markdown("---")
            st.markdown("#### 🎯 Roteiro e Bipe de Caixas (Antierros)")
            
            todos_separados = True
            for i, obj in enumerate(itens_ordenados):
                item = obj['item_original']
                vinho_est = obj['vinho_estoque']
                loc = obj['localizacao']
                
                status_cor = "🟢" if item.get('separado') else "⏳"
                if not item.get('separado'): todos_separados = False
                
                key_bip_state = f"bip_val_{idx_ped}_{i}"
                if key_bip_state not in st.session_state.codigos_bipados_conferencia:
                    st.session_state.codigos_bipados_conferencia[key_bip_state] = ""

                with st.expander(f"{status_cor} {item['nome']} {item.get('safra', '')} | Qtd: {item['quantidade']} caixas | 📍 {loc}"):
                    if vinho_est:
                        st.markdown(f"**Estoque:** Safra no Galpão: **{vinho_est.get('safra')}** | Local: **{vinho_est.get('localizacao')} - Lado: {vinho_est.get('lado')}** | C. Barras Cadastrado: `{vinho_est.get('codigo_barras', 'Não cadastrado')}`")
                        
                        safra_matriz = str(item.get('safra', '')).strip()
                        safra_galpao = str(vinho_est.get('safra', '')).strip()
                        
                        if safra_matriz and safra_galpao and safra_matriz != safra_galpao:
                            st.warning(f"⚠️ **Alerta de Safra Divergente!** Pedido pediu safra `{safra_matriz}`, mas no galpão é `{safra_galpao}`.")
                        
                        st.markdown("---")
                        st.markdown("📲 **Conferência Física:**")
                        qtd_informada = st.number_input("Quantidade que está levando:", min_value=1, value=item.get('quantidade', 1), key=f"qtd_inf_{idx_ped}_{i}")
                        
                        if qtd_informada != item['quantidade']:
                            st.error(f"❌ **Divergência de Quantidade!** O pedido pedia **{item['quantidade']}**, mas você informou **{qtd_informada}**.")
                        else:
                            st.success("✅ Quantidade confere com o pedido.")
                            
                        bip_caixa = st.text_input("Código de barras da caixa (ou escaneie abaixo):", value=st.session_state.codigos_bipados_conferencia[key_bip_state], key=f"bip_txt_{idx_ped}_{i}")
                        if bip_caixa != st.session_state.codigos_bipados_conferencia[key_bip_state]:
                            st.session_state.codigos_bipados_conferencia[key_bip_state] = bip_caixa

                        st.markdown("📷 **Leitor contínuo (Estilo App de Banco):**")
                        componente_leitor_barcode(key_bip_state)

                        codigo_atual_conferido = st.session_state.codigos_bipados_conferencia[key_bip_state]
                        if codigo_atual_conferido:
                            cb_cadastrado = str(vinho_est.get('codigo_barras', '')).strip()
                            if cb_cadastrado and codigo_atual_conferido.strip() == cb_cadastrado:
                                st.success("✅ Código de barras conferido e validado com sucesso!")
                            else:
                                st.error("❌ **Erro:** Este código de barras NÃO corresponde ao vinho da lista!")
                        
                        if st.button(f"✅ Confirmar Separação do Item #{i+1}", key=f"btn_sep_{idx_ped}_{i}"):
                            cb_cadastrado = str(vinho_est.get('codigo_barras', '')).strip()
                            if codigo_atual_conferido.strip() != cb_cadastrado:
                                st.error("Não é possível concluir: o código de barras bipado está incorreto ou diverge do vinho!")
                            elif qtd_informada != item['quantidade']:
                                st.error("Não é possível concluir: a quantidade informada diverge do pedido.")
                            else:
                                item['separado'] = True
                                item['qtd_separada'] = qtd_informada
                                salvar_pedidos(st.session_state.pedidos)
                                registrar_log(st.session_state.usuario_logado['nome'], "Separar Item Pedido", f"{item['nome']} (Qtd: {qtd_informada})")
                                st.success("Item confirmado e separado!")
                                st.rerun()
                    else:
                        st.error("❌ Vinho não encontrado no estoque do galpão.")

            st.markdown("---")
            if todos_separados:
                st.balloons()
                st.success("🎉 Pedido 100% separado e conferido sem divergências!")
                
                if st.button("💾 Finalizar e Arquivar Pedido", use_container_width=True):
                    pedido_atual['status'] = "Concluído"
                    salvar_pedidos(st.session_state.pedidos)
                    registrar_log(st.session_state.usuario_logado['nome'], "Concluir Pedido", pedido_atual['id'])
                    st.success("Pedido finalizado com sucesso!")
                    st.rerun()
                
                texto_resumo = f"📦 *Relatório de Separação - {pedido_atual['id']}*\nStatus: Concluído ✅\nSeparado por: {st.session_state.usuario_logado['nome']}\n\n*Itens Conferidos:*\n"
                for obj in itens_ordenados:
                    it = obj['item_original']
                    texto_resumo += f"- {it['nome']} ({it.get('safra','')}) | Qtd: {it['quantidade']} | 📍 {obj['localizacao']}\n"
                
                url_wapp = f"https://api.whatsapp.com/send?text={urllib.parse.quote(texto_resumo)}"
                
                st.markdown(f"""
                    <div style="margin-top: 10px;">
                        <a href="{url_wapp}" target="_blank" style="
                            background-color: #25D366; 
                            color: white; 
                            padding: 12px 16px; 
                            border-radius: 12px; 
                            text-decoration: none; 
                            font-weight: 600; 
                            width: 100%; 
                            display: block; 
                            text-align: center; 
                            font-family: 'Poppins', sans-serif; 
                            font-size: 1rem;
                            box-shadow: 0 4px 6px rgba(0,0,0,0.1);">
                            📤 Compartilhar via WhatsApp
                        </a>
                    </div>
                """, unsafe_allow_html=True)

elif st.session_state.menu_atual == "Filtros":
    st.subheader("🔍 Busca por Local ou Nome")
    termo_digitado = st.text_input("Filtrar por nome ou corredor/pallet:", value=st.session_state.termo_busca)
    termo = termo_digitado.strip().title()
    
    tp = termo.lower() if termo else st.session_state.termo_busca.lower()
    if tp:
        res = [v for v in st.session_state.estoque if tp in v.get("nome", "").lower() or tp in v.get("localizacao", "").lower() or tp in v.get("lado", "").lower()]
        if res:
            for v in sorted(res, key=lambda x: x.get("nome", "").lower()):
                img_tag = f"<br><img src='app/static/{v.get('foto')}' width='100' style='border-radius: 8px; margin-top: 8px;'>" if v.get('foto') and os.path.exists(os.path.join(PASTA_FOTOS, v.get('foto'))) else ""
                st.markdown(f"<div class='wine-card'><div class='wine-title'>🍷 {v.get('nome')} ({v.get('safra')})</div><p>Tipo: <b>{v.get('tipo', 'N/A')}</b> | Caixa: <b>{v.get('caixa', 'N/A')}</b><br><span class='badge-pallet-grande'>📍 {v.get('localizacao')} - Lado: {v.get('lado', 'N/A')}</span>{img_tag}</p></div>", unsafe_allow_html=True)
        else:
            st.warning("Nenhum vinho encontrado para este filtro/local.")

elif st.session_state.menu_atual == "MapaSeparacao":
    st.subheader("🗺️ Mapa de Separação")
    arq = st.file_uploader("Envie o arquivo com a lista de vinhos", type=["xlsx", "xls", "txt"])
    txt_man_input = st.text_area("Ou cole a lista de vinhos manualmente aqui:")
    txt_man = txt_man_input.title() if txt_man_input else ""
    
    if st.button("Gerar Rota / Mapa"):
        linhas = extrair_linhas_de_arquivo(arq) if arq else []
        if txt_man.strip():
            linhas.extend([l.strip().title() for l in txt_man.split("\n") if l.strip()])
            
        if linhas:
            encontrados = []
            for v in st.session_state.estoque:
                nome_vinho = v.get("nome", "").lower().strip()
                if any(l.lower().strip() in nome_vinho or nome_vinho in l.lower().strip() for l in linhas if len(l.strip()) > 2):
                    if v not in encontrados:
                        encontrados.append(v)
            
            if encontrados:
                st.success(f"Foram encontrados {len(encontrados)} vinhos para separação:")
                for v in sorted(encontrados, key=lambda x: x.get("nome", "").lower()):
                    st.markdown(f"<div class='wine-card'><div class='wine-title'>🍷 {v.get('nome')} ({v.get('safra', 'N/A')})</div><p>Tipo: <b>{v.get('tipo', 'N/A')}</b> | Caixa: <b>{v.get('caixa', 'N/A')}</b><br><span class='badge-pallet-grande'>📍 {v.get('localizacao', 'N/A')} - Lado: {v.get('lado', 'N/A')}</span></p></div>", unsafe_allow_html=True)
            else:
                st.warning("Nenhum vinho do arquivo corresponde aos cadastrados no estoque.")
        else:
            st.error("Envie um arquivo ou digite pelo menos um nome.")

elif st.session_state.menu_atual == "Scanner":
    st.subheader("📷 Escanear QR Code do Local")
    foto = st.camera_input("Capturar Foto do QR Code do Pallet/Prateleira")
    if foto and OPENCV_DISPONIVEL:
        img = cv2.imdecode(np.frombuffer(foto.getvalue(), np.uint8), cv2.IMREAD_COLOR)
        val, _, _ = cv2.QRCodeDetector().detectAndDecode(img)
        if val:
            termo_lido = val.strip().lower()
            st.success(f"Local Identificado: {val}")
            st.markdown("### 🍷 Vinhos neste local:")
            resultados = [v for v in st.session_state.estoque if termo_lido in v.get("localizacao", "").lower()]
            if resultados:
                for v in sorted(resultados, key=lambda x: x.get("nome", "").lower()):
                    st.markdown(f"<div class='wine-card'><div class='wine-title'>🍷 {v.get('nome')} ({v.get('safra', 'N/A')})</div><p>Tipo: <b>{v.get('tipo', 'N/A')}</b> | Caixa: <b>{v.get('caixa', 'N/A')}</b><br><span class='badge-pallet-grande'>📍 {v.get('localizacao', 'N/A')} - Lado: {v.get('lado', 'N/A')}</span></p></div>", unsafe_allow_html=True)
            else:
                st.warning("Nenhum vinho cadastrado neste exato local.")
        else:
            st.warning("Nenhum QR Code detectado na imagem. Tente novamente.")
    elif not OPENCV_DISPONIVEL:
        st.error("A biblioteca OpenCV não está disponível para leitura de QR Code por câmera neste ambiente.")

elif st.session_state.menu_atual == "Estoque":
    st.subheader("🍷 Estoque Completo do Galpão")
    if st.session_state.estoque:
        col_exp1, col_exp2 = st.columns(2)
        with col_exp1:
            df_estoque = pd.DataFrame(st.session_state.estoque)
            csv = df_estoque.to_csv(index=False).encode('utf-8')
            st.download_button("📥 Baixar CSV do Estoque", data=csv, file_name="estoque_galpao.csv", mime="text/csv", use_container_width=True)
        with col_exp2:
            if st.button("🗑️ Apagar Todo o Estoque (Cuidado)", use_container_width=True):
                if st.session_state.usuario_logado.get('cargo') in ["Administrador", "Desenvolvedor"]:
                    st.session_state.estoque = []
                    salvar_dados([])
                    registrar_log(st.session_state.usuario_logado['nome'], "Apagar Estoque", "Limpeza total do estoque")
                    st.success("Estoque limpo com sucesso!")
                    st.rerun()
                else:
                    st.error("Apenas administradores podem apagar o estoque.")

        st.markdown("---")
        for v in st.session_state.estoque:
            st.markdown(f"<div class='wine-card'><div class='wine-title'>🍷 {v.get('nome')} ({v.get('safra')})</div><p>Tipo: <b>{v.get('tipo')}</b> | Caixa: <b>{v.get('caixa')}</b><br><span class='badge-pallet-grande'>📍 {v.get('localizacao')} - Lado: {v.get('lado')}</span> | C. Barras: <code>{v.get('codigo_barras', 'N/A')}</code></p></div>", unsafe_allow_html=True)
    else:
        st.info("Nenhum vinho cadastrado no estoque.")

elif st.session_state.menu_atual == "Cadastrar":
    st.subheader("➕ Cadastrar Novo Vinho")
    with st.form("form_cadastro_vinho"):
        nome_cad = st.text_input("Nome do Vinho *").strip().title()
        tipo_cad = st.selectbox("Tipo de Vinho", ["Tinto", "Branco", "Rosé", "Espumante", "Fortificado / Doce"])
        safra_cad = st.text_input("Safra (Ano)").strip()
        
        c_corredor = st.selectbox("Corredor", LISTA_CORREDORES)
        c_tipo_local = st.selectbox("Tipo de Local", LISTA_LOCAIS_TIPO)
        c_num_local = st.selectbox("Número do Local", LISTA_NUMEROS_LOCAL)
        localizacao_cad = f"{c_corredor} - {c_tipo_local} {c_num_local}"
        
        lado_cad = st.selectbox("Lado do Corredor", LISTA_LADOS)
        caixa_cad = st.selectbox("Embalagem", OPCOES_CAIXA)
        
        codigo_barras_cad = st.text_input("Código de Barras da Caixa (Bipe ou Digite)", value=st.session_state.codigo_capturado_cadastro).strip()
        
        if st.form_submit_button("Salvar Novo Vinho"):
            if nome_cad:
                novo_vinho = {
                    "nome": nome_cad,
                    "tipo": tipo_cad,
                    "safra": safra_cad,
                    "localizacao": localizacao_cad,
                    "lado": lado_cad,
                    "caixa": caixa_cad,
                    "codigo_barras": codigo_barras_cad,
                    "foto": ""
                }
                st.session_state.estoque.append(novo_vinho)
                salvar_dados(st.session_state.estoque)
                registrar_log(st.session_state.usuario_logado['nome'], "Cadastrar Vinho", f"{nome_cad} ({safra_cad}) em {localizacao_cad}")
                st.session_state.codigo_capturado_cadastro = ""
                st.success("Vinho cadastrado com sucesso!")
                st.rerun()
            else:
                st.error("O campo 'Nome do Vinho' é obrigatório.")

    st.markdown("---")
    st.markdown("📷 **Leitor de Código de Barras para Cadastro:**")
    componente_leitor_barcode("cad_barcode")

elif st.session_state.menu_atual == "GerarQR":
    st.subheader("📱 Gerar QR Code para Impressão")
    tipo_qr = st.radio("O que deseja gerar?", ["Localização (Corredor/Pallet)", "Vinho Específico"], horizontal=True)
    
    if "Local" in tipo_qr:
        c_qr_corr = st.selectbox("Corredor", LISTA_CORREDORES, key="qr_corr")
        c_qr_tipo = st.selectbox("Tipo", LISTA_LOCAIS_TIPO, key="qr_tipo")
        c_qr_num = st.selectbox("Número", LISTA_NUMEROS_LOCAL, key="qr_num")
        texto_gerar = f"{c_qr_corr} - {c_qr_tipo} {c_qr_num}"
    else:
        if st.session_state.estoque:
            nomes_v = [f"{v['nome']} ({v.get('safra', '')}) - {v.get('localizacao', '')}" for v in st.session_state.estoque]
            escolha_v_qr = st.selectbox("Selecione o Vinho:", nomes_v)
            texto_gerar = escolha_v_qr
        else:
            texto_gerar = "Nenhum vinho"
            
    st.markdown(f"### QR Code para: `{texto_gerar}`")
    url_img = gerar_qr_code_api(texto_gerar)
    st.image(url_img, width=250)
    st.markdown(f"[🔗 Abrir imagem do QR Code em nova aba]({url_img})", unsafe_allow_html=True)

elif st.session_state.menu_atual == "Editar":
    st.subheader("✏️ Editar ou Excluir Vinho")
    if st.session_state.estoque:
        nomes_estoque = [f"{v['nome']} ({v.get('safra', 'N/A')}) - 📍 {v.get('localizacao', 'N/A')}" for v in st.session_state.estoque]
        vinho_escolhido_str = st.selectbox("Selecione o vinho para editar:", nomes_estoque)
        idx_vinho = nomes_estoque.index(vinho_escolhido_str)
        v_atual = st.session_state.estoque[idx_vinho]
        
        with st.form("form_edicao_vinho"):
            n_nome = st.text_input("Nome do Vinho", value=v_atual.get('nome', '')).strip().title()
            n_tipo = st.selectbox("Tipo", ["Tinto", "Branco", "Rosé", "Espumante", "Fortificado / Doce"], index=["Tinto", "Branco", "Rosé", "Espumante", "Fortificado / Doce"].index(v_atual.get('tipo', 'Tinto')) if v_atual.get('tipo') in ["Tinto", "Branco", "Rosé", "Espumante", "Fortificado / Doce"] else 0)
            n_safra = st.text_input("Safra", value=v_atual.get('safra', '')).strip()
            n_local = st.text_input("Localização", value=v_atual.get('localizacao', '')).strip()
            n_lado = st.selectbox("Lado", LISTA_LADOS, index=LISTA_LADOS.index(v_atual.get('lado', 'Direito')) if v_atual.get('lado') in LISTA_LADOS else 0)
            n_caixa = st.selectbox("Embalagem", OPCOES_CAIXA, index=OPCOES_CAIXA.index(v_atual.get('caixa', OPCOES_CAIXA[0])) if v_atual.get('caixa') in OPCOES_CAIXA else 0)
            n_cb = st.text_input("Código de Barras", value=v_atual.get('codigo_barras', '')).strip()
            
            col_b1, col_b2 = st.columns(2)
            with col_b1:
                salvar_edicao = st.form_submit_button("💾 Salvar Alterações", use_container_width=True)
            with col_b2:
                excluir_vinho = st.form_submit_button("🗑️ Excluir Vinho", use_container_width=True)
                
            if salvar_edicao:
                st.session_state.estoque[idx_vinho] = {
                    "nome": n_nome,
                    "tipo": n_tipo,
                    "safra": n_safra,
                    "localizacao": n_local,
                    "lado": n_lado,
                    "caixa": n_caixa,
                    "codigo_barras": n_cb,
                    "foto": v_atual.get('foto', '')
                }
                salvar_dados(st.session_state.estoque)
                registrar_log(st.session_state.usuario_logado['nome'], "Editar Vinho", f"{n_nome} ({n_safra})")
                st.success("Alterações salvas com sucesso!")
                st.rerun()
                
            if excluir_vinho:
                vinho_removido = st.session_state.estoque.pop(idx_vinho)
                salvar_dados(st.session_state.estoque)
                registrar_log(st.session_state.usuario_logado['nome'], "Excluir Vinho", f"{vinho_removido.get('nome')}")
                st.success("Vinho excluído com sucesso!")
                st.rerun()
    else:
        st.info("Nenhum vinho cadastrado para editar.")

elif st.session_state.menu_atual == "Historico":
    st.subheader("📋 Histórico de Auditoria e Atividades")
    logs = carregar_logs()
    if logs:
        for l in logs[:50]:
            st.markdown(f"<div class='wine-card'><p><b>🕒 {l.get('data_hora')}</b> | Usuário: <b>{l.get('usuario')}</b><br>Ação: <b>{l.get('acao')}</b> | Detalhes: {l.get('detalhes')}</p></div>", unsafe_allow_html=True)
    else:
        st.info("Nenhum registro no histórico.")

elif st.session_state.menu_atual == "GerenciarUsuarios":
    st.subheader("⚙️ Gerenciamento de Contas de Usuários")
    if st.session_state.usuario_logado.get('cargo') == "Desenvolvedor":
        if st.session_state.usuarios:
            for i, u in enumerate(st.session_state.usuarios):
                st.markdown(f"<div class='wine-card'><p>👤 Nome: <b>{u.get('nome')}</b> | Cargo: <b>{u.get('cargo', 'Operador')}</b></p></div>", unsafe_allow_html=True)
        else:
            st.info("Nenhum usuário cadastrado.")
    else:
        st.error("Acesso restrito ao desenvolvedor.")
